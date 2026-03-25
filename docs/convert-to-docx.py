from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Title
title = doc.add_heading('Calibração da IA — Avaliação de Desempenho Seazone', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Meta info
meta = doc.add_paragraph()
meta.add_run('Para: ').bold = True
meta.add_run('CEO')
meta.add_run('\nObjetivo: ').bold = True
meta.add_run('Alinhar como a IA deve interpretar cada critério de avaliação, considerando o nível do cargo e o que significa cada nota (1 a 5) na prática da Seazone.')

doc.add_paragraph()
how = doc.add_paragraph()
how.add_run('Como preencher: ').bold = True
how.add_run('Para cada critério, descreva com suas palavras o que você espera ver em cada nota, considerando os diferentes níveis de cargo. Pode usar exemplos reais (sem nomes) ou hipotéticos. Quanto mais específico, melhor a IA vai calibrar.')

levels_text = doc.add_paragraph()
levels_text.add_run('Níveis de cargo: ').bold = True
levels_text.add_run('Estágio | Analista | Especialista | Coordenador | Gerente | Diretor | C-Level')

doc.add_page_break()

# Questions data
questions = [
    ("1", "Sangue no Olho", "Com que frequência essa pessoa demonstra Sangue no Olho ao enfrentar obstáculos e sustentar esforço até a conclusão dos desafios?"),
    ("2", "Atitude de Dono", "Com que frequência essa pessoa demonstra Atitude de Dono nas decisões e responsabilidades que lhe cabem?"),
    ("3", "Foco em Fatos e Dados", "Com que frequência essa pessoa utiliza Foco em Fatos e Dados para fundamentar decisões e argumentos?"),
    ("4", "Priorize e Simplifique", "Com que frequência essa pessoa demonstra capacidade de Priorizar e Simplificar?"),
    ("5", "Escopo da Função", "Com que frequência essa pessoa atua de forma madura e adequada ao Escopo da Função?"),
    ("6", "Entregas de Valor", "Com que frequência essa pessoa realiza Entregas de Valor relevantes?"),
    ("7", "Consistência", "Com que frequência essa pessoa mantém Consistência no desempenho ao longo do tempo?"),
    ("8", "Pensar Fora da Caixa", "Com que frequência essa pessoa demonstra capacidade de Pensar Fora da Caixa ao lidar com desafios?"),
    ("9", "Organização", "Com que frequência essa pessoa demonstra Organização na gestão das suas atividades?"),
    ("10", "Adaptabilidade", "Com que frequência essa pessoa demonstra Adaptabilidade diante de mudanças?"),
    ("11", "Comunicação", "Com que frequência essa pessoa demonstra Comunicação clara e adequada ao seu contexto de atuação?"),
    ("12", "Colaboração", "Com que frequência essa pessoa demonstra Colaboração positiva?"),
    ("13", "Uso de IA", "Com que frequência essa pessoa utiliza Inteligência Artificial para potencializar seu trabalho e resultados?"),
]

levels = ["Estágio", "Analista", "Especialista", "Coordenador", "Gerente", "Diretor", "C-Level"]

notas = [
    ("Nota 1 (Insuficiente)", "Descreva um comportamento ou situação que seria nota 1. O que a pessoa faz (ou deixa de fazer)?"),
    ("Nota 2 (Abaixo do esperado)", "O que diferencia alguém nota 2 de nota 1? Qual o comportamento típico?"),
    ("Nota 3 (Dentro do esperado)", "Esse é o padrão — a maioria deveria ser nota 3. O que é \"o esperado\" nesse critério?"),
    ("Nota 4 (Acima do esperado)", "O que alguém precisa fazer ALÉM do esperado para merecer nota 4? O que separa o 3 do 4?"),
    ("Nota 5 (Excepcional)", "Nota 5 é rara. Descreva o que seria excepcional. Tem algum exemplo (real ou hipotético)?"),
]

def add_response_box(doc, lines=3):
    """Add blank lines for response"""
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run('_' * 80)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.size = Pt(9)

for num, title_text, description in questions:
    # Section heading
    doc.add_heading(f'{num}. {title_text}', level=1)

    # Question description
    p = doc.add_paragraph()
    p.add_run('Pergunta na avaliação: ').bold = True
    run = p.add_run(f'"{description}"')
    run.italic = True

    # 1. Meaning
    doc.add_heading(f'{num}.1. O que "{title_text}" significa para você na Seazone?', level=2)
    add_response_box(doc, 3)

    # 2. By level table
    doc.add_heading(f'{num}.2. Como esse valor se manifesta em cada nível?', level=2)

    table = doc.add_table(rows=len(levels) + 1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0]
    hdr.cells[0].text = 'Nível'
    hdr.cells[1].text = 'O que você espera ver nesse nível?'
    for cell in hdr.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.3)
        row.cells[1].width = Inches(5.2)

    for i, level in enumerate(levels):
        row = table.rows[i + 1]
        row.cells[0].text = level
        row.cells[1].text = ''
        # Add minimum height for writing
        for paragraph in row.cells[1].paragraphs:
            paragraph.paragraph_format.space_after = Pt(20)

    doc.add_paragraph()  # spacing

    # 3. Examples by score
    doc.add_heading(f'{num}.3. Exemplos por nota', level=2)

    for nota_label, nota_desc in notas:
        p = doc.add_paragraph()
        p.add_run(f'{nota_label}: ').bold = True
        p.add_run(nota_desc)
        add_response_box(doc, 2)
        doc.add_paragraph()  # spacing

    # 4. Common mistakes
    doc.add_heading(f'{num}.4. Erros comuns que a IA deve evitar nesse critério:', level=2)
    add_response_box(doc, 2)

    doc.add_page_break()

# General calibration questions
doc.add_heading('Perguntas Gerais de Calibração', level=1)

general_questions = [
    ("A. Distribuição esperada de notas",
     "Na sua visão, qual deveria ser a distribuição aproximada de notas na empresa?\n(Ex: \"80% nota 3, 10% nota 4, 5% nota 2, 3% nota 5, 2% nota 1\")"),
    ("B. Diferença entre níveis de cargo",
     "Um analista nota 4 e um coordenador nota 3 podem ter comportamentos parecidos? Como a IA deve lidar com isso?"),
    ("C. Antiguidade vs. desempenho",
     "Alguém com muitos anos de empresa mas que \"faz o básico\" deveria ser nota 3 ou abaixo? A senioridade influencia?"),
    ("D. Autoavaliação",
     "Autoavaliações tendem a ser infladas. Qual o nível de desconto que a IA deve aplicar? Deve questionar mais?"),
    ("E. Avaliação de pares",
     "Pares podem ter visibilidade parcial. A IA deve ser mais flexível com pares que dizem \"não tenho visibilidade sobre isso\"?"),
    ("F. Situações ambíguas",
     "Tem alguma situação que você já viu ser mal avaliada? Algo que parece nota 4 mas na real é 3 (ou vice-versa)?"),
    ("G. Palavras ou expressões de alerta",
     "Quais expressões devem acender um alerta na IA?\n(Ex: \"sempre dá o seu melhor\" sem exemplo concreto = nota máxima 3)"),
]

for gtitle, gdesc in general_questions:
    doc.add_heading(gtitle, level=2)
    doc.add_paragraph(gdesc)
    add_response_box(doc, 3)
    doc.add_paragraph()

# Footer
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Obrigado por preencher! Suas respostas serão usadas para calibrar a IA da plataforma IAVD, garantindo que as notas reflitam a cultura e os padrões da Seazone.')
run.bold = True
run.font.size = Pt(10)

# Save
output_path = r'C:\Users\compu\projetos\projeto-iavd\docs\Calibracao-IA-CEO.docx'
doc.save(output_path)
print(f'Documento salvo em: {output_path}')
